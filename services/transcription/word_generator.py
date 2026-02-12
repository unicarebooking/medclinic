"""
Word document generator for transcriptions
Creates RTL Hebrew Word documents with transcription content
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path


def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL direction"""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)


def generate_word_document(
    transcription_text: str,
    doctor_name: str = "",
    doctor_specialization: str = "",
    patient_name: str = "",
    original_filename: str = "",
    duration_seconds: float = 0,
    output_path: str = None,
) -> str:
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'David'
    font.size = Pt(12)

    # Set RTL for the entire document
    section = doc.sections[0]
    sectPr = section._sectPr
    bidi = sectPr.makeelement(qn('w:bidi'), {})
    sectPr.append(bidi)

    # --- Header ---
    header_para = doc.add_paragraph()
    set_rtl_paragraph(header_para)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run('DOCTOR SEARCH')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(59, 130, 246)

    subtitle = doc.add_paragraph()
    set_rtl_paragraph(subtitle)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('תמלול שיחה רפואית')
    run.bold = True
    run.font.size = Pt(16)

    # --- Separator ---
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('─' * 50)
    run.font.color.rgb = RGBColor(156, 163, 175)

    # --- Details ---
    now = datetime.now()
    date_str = now.strftime('%d/%m/%Y')
    time_str = now.strftime('%H:%M')

    details = [
        ('תאריך', date_str),
        ('שעה', time_str),
    ]

    if doctor_name:
        details.insert(0, ('רופא/ה', f'ד"ר {doctor_name}'))
    if doctor_specialization:
        details.insert(1, ('התמחות', doctor_specialization))
    if patient_name:
        details.append(('מטופל/ת', patient_name))
    if original_filename:
        details.append(('קובץ מקור', original_filename))
    if duration_seconds > 0:
        minutes = int(duration_seconds // 60)
        seconds = int(duration_seconds % 60)
        details.append(('משך ההקלטה', f'{minutes} דקות ו-{seconds} שניות'))

    for label, value in details:
        p = doc.add_paragraph()
        set_rtl_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_label = p.add_run(f'{label}: ')
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_value = p.add_run(value)
        run_value.font.size = Pt(11)

    # --- Separator ---
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep2.add_run('─' * 50)
    run.font.color.rgb = RGBColor(156, 163, 175)

    # --- Transcription Title ---
    title = doc.add_paragraph()
    set_rtl_paragraph(title)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = title.add_run('תוכן התמלול:')
    run.bold = True
    run.font.size = Pt(14)

    # --- Transcription Content ---
    content = doc.add_paragraph()
    set_rtl_paragraph(content)
    content.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = content.add_run(transcription_text)
    run.font.size = Pt(12)
    paragraph_format = content.paragraph_format
    paragraph_format.line_spacing = 1.5

    # --- Footer ---
    doc.add_paragraph()
    sep3 = doc.add_paragraph()
    sep3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep3.add_run('─' * 50)
    run.font.color.rgb = RGBColor(156, 163, 175)

    footer = doc.add_paragraph()
    set_rtl_paragraph(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('מסמך זה נוצר באופן אוטומטי על ידי מערכת DOCTOR SEARCH')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(156, 163, 175)

    privacy = doc.add_paragraph()
    set_rtl_paragraph(privacy)
    privacy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = privacy.add_run('מסמך חסוי - לשימוש רפואי בלבד')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(220, 38, 38)

    # Save
    if output_path is None:
        output_dir = Path(__file__).parent / 'outputs'
        output_dir.mkdir(exist_ok=True)
        timestamp = now.strftime('%Y%m%d_%H%M%S')
        output_path = str(output_dir / f'transcription_{timestamp}.docx')

    doc.save(output_path)
    return output_path
